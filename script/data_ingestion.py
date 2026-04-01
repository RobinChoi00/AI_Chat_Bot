import pandas as pd
import logging
import docx 
import os
import glob
from dotenv import load_dotenv

from langchain_openai import OpenAIEmbeddings
from langchain_community.vectorstores import FAISS
from langchain_core.documents import Document

logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

load_dotenv()

# 🔥 [아키텍처 경로 동기화]
BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__))) 
DATA_DIR = os.path.join(BASE_DIR, "data")          
RAW_DIR = os.path.join(BASE_DIR, "raw_data")       
FAISS_DIR = os.path.join(BASE_DIR, "faiss_index")  

def extract_text_from_docx(file_path):
    """Word 문서에서 텍스트를 추출하고 의미 단위로 청킹(Chunking)합니다."""
    doc = docx.Document(file_path)
    chunks = []
    current_chunk = ""
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text: continue
        current_chunk += text + " "
        if len(current_chunk) > 400: 
            chunks.append(current_chunk.strip())
            current_chunk = ""
    if current_chunk:
        chunks.append(current_chunk.strip())
    return chunks

def build_products_index():
    logger.info("🚀 [통합 아키텍처] 다중 정책 문서 & 메타데이터 적용 Vector DB 굽기 가동")
    
    os.makedirs(FAISS_DIR, exist_ok=True)
    os.makedirs(DATA_DIR, exist_ok=True) 

    documents = []

    # 1. 💡 [Staging Layer] 메타데이터를 포함한 정제된 상품 데이터 섭취
    clean_csv = os.path.join(DATA_DIR, "cleaned_osaki_products.csv")
    if os.path.exists(clean_csv):
        logger.info(f"📊 정제된 상품 데이터 섭취 중: {os.path.basename(clean_csv)}")
        try:
            df = pd.read_csv(clean_csv)
            
            # 💡 [방어적 프로그래밍] 결측치(NaN)로 인한 타입 에러 원천 차단
            df['price'] = df.get('price', 0.0).fillna(0.0)
            df['product_type'] = df.get('product_type', 'Unknown').fillna('Unknown')
            df['vendor'] = df.get('vendor', 'Unknown').fillna('Unknown')

            for _, row in df.iterrows():
                content = str(row['content'])
                
                # 💡 [핵심] 풍부해진 메타데이터 주입 (Self-Querying 필터링 기반 마련)
                metadata = {
                    "source": str(row['source']), 
                    "type": "product",
                    "price": float(row['price']),
                    "product_type": str(row['product_type']),
                    "vendor": str(row['vendor'])
                }
                documents.append(Document(page_content=content, metadata=metadata))
        except Exception as e:
            logger.error(f"🚨 CSV 처리 실패: {e}")
    else:
        logger.error(f"🚨 {clean_csv} 파일이 없습니다. 파이프라인을 중단합니다.")
        return 

    # 2. 💡 [Raw Layer] 하드코딩 제거: raw_data 내의 모든 .docx 자동 섭취
    policy_files = glob.glob(os.path.join(RAW_DIR, "*.docx"))
    
    if policy_files:
        for file_path in policy_files:
            file_name = os.path.basename(file_path)
            logger.info(f"📄 정책 문서 섭취 중: {file_name}")
            try:
                chunks = extract_text_from_docx(file_path)
                for chunk in chunks:
                    # 💡 파일명 자체를 출처(source)로 사용하여 메타데이터 생성
                    metadata = {"source": file_name, "type": "policy"}
                    documents.append(Document(page_content=chunk, metadata=metadata))
            except Exception as e:
                logger.error(f"🚨 {file_name} 문서 처리 실패: {e}")
    else:
        logger.warning(f"⚠️ {RAW_DIR} 폴더 내에 정책 문서(.docx)를 찾을 수 없습니다.")

    if not documents:
        logger.error("🚨 인덱싱할 데이터가 존재하지 않습니다.")
        return

    # 3. Load: 벡터화 및 저장
    logger.info(f"🧠 총 {len(documents)}개의 문서를 OpenAI 임베딩으로 변환 중...")
    try:
        embeddings = OpenAIEmbeddings() 
        vector_db = FAISS.from_documents(documents, embeddings)

        save_path = os.path.join(FAISS_DIR, "osaki_products")
        vector_db.save_local(save_path)
        logger.info(f"🎉 성공! 확장된 메타데이터와 다중 정책이 포함된 DB가 '{save_path}'에 구축되었습니다.")
        
    except Exception as e:
        logger.error(f"🚨 벡터 DB 생성 중 치명적 에러: {e}")

if __name__ == "__main__":
    build_products_index()