import os
import sys
import glob
import json
import pandas as pd
import docx
from dotenv import load_dotenv
from langchain_openai import OpenAIEmbeddings
from langchain_community.vectorstores import FAISS as LC_FAISS
from langchain_core.documents import Document

# 💡 [아키텍처] 마스터 파이프라인 경로 설정
load_dotenv(override=True)
BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
RAW_DIR = os.path.join(BASE_DIR, "raw_data")
DATA_DIR = os.path.join(BASE_DIR, "data")
FAISS_DIR = os.path.join(BASE_DIR, "faiss_index")

# Pull the configured embedding model from the central config so re-builds
# stay in lock-step with what the chat server is querying at runtime.
sys.path.insert(0, BASE_DIR)
try:
    from config import EMBEDDING_MODEL
except Exception:
    EMBEDDING_MODEL = os.environ.get("OPENAI_EMBEDDING_MODEL", "text-embedding-3-small")


class MasterIngester:
    def __init__(self):
        # 💰 text-embedding-3-small: ~5x cheaper than ada-002 with comparable
        # or better recall. If you flip back to ada-002 you MUST rebuild this
        # index since the vector dimensions differ.
        print(f"🧠 Embedding model: {EMBEDDING_MODEL}")
        self.embeddings = OpenAIEmbeddings(model=EMBEDDING_MODEL)
        # 💡 [리팩토링 핵심] 도메인별 3개의 바구니(딕셔너리) 생성
        self.domain_docs = {
            "freshdesk_qa": [],   
            "web_data": [],       
            "osaki_products": []  
        }

    # ==========================================
    # 🧠 1번 뇌: freshdesk_qa (CS & 에러 전용)
    # ==========================================
    def process_error_manuals(self):
        print("🔍 [1/7] 에러 매뉴얼 CSV 파싱 중... -> [freshdesk_qa] 할당")
        for file_name, skip_rows in {"Auto-Check.csv": 1, "fault_judgment.csv": 5}.items():
            file_path = os.path.join(RAW_DIR, file_name)
            if not os.path.exists(file_path): continue
            
            # 💡 [아키텍처] 판다스의 강제 형변환(Float)을 막고 원본 텍스트(63) 그대로 보존합니다.
            df = pd.read_csv(file_path, skiprows=skip_rows, encoding='utf-8-sig', dtype=str).fillna("")
            
            # 💡 [아키텍처] 대소문자 무시 & 앞뒤 공백 제거로 더러운 헤더를 깔끔하게 평탄화(Normalization)
            df.columns = df.columns.str.strip().str.lower()
            
            for _, row in df.iterrows():
                # 번호 추출 (No. 또는 Code No.)
                code_no = str(row.get("code no.", row.get("no.", ""))).strip()
                if not code_no or code_no.lower() == 'nan': continue
                
                # 💡 [핵심] 컬럼명이 달라도 모두 잡아내는 동적 매핑 로직
                symptom = str(row.get("phenomenon", row.get("problem description", ""))).strip()
                troubleshooting = str(row.get("troubleshooting steps", row.get("steps of shooting the trouble", ""))).strip()
                
                # 빈 데이터 필터링 방어 로직
                if not symptom and not troubleshooting: continue
                
                content = f"[Error Code]: {code_no}\n[Symptom]: {symptom}\n[Troubleshooting]: {troubleshooting}"
                self.domain_docs["freshdesk_qa"].append(Document(page_content=content, metadata={"type": "error_code", "error_code": code_no}))

    def process_qa_reports(self):
        print("🔍 [2/7] CS Q&A CSV 파싱 중... -> [freshdesk_qa] 할당")
        file_path = os.path.join(RAW_DIR, "Warranty Daily Report - Q&A.csv")
        if os.path.exists(file_path):
            df = pd.read_csv(file_path, skiprows=1).fillna("")
            for _, row in df.iterrows():
                issue = str(row.iloc[0]).strip()
                if not issue or issue == "N/A": continue
                content = f"[Issue]: {issue}\n[Diagnostic]: {row.iloc[1]}\n[Solution]: {row.iloc[2]}"
                self.domain_docs["freshdesk_qa"].append(Document(page_content=content, metadata={"type": "troubleshooting"}))

    def process_freshdesk_tickets(self):
        print("🔍 [3/7] Freshdesk 고객 상담 티켓 파싱 중... -> [freshdesk_qa] 할당")
        file_path = os.path.join(DATA_DIR, "freshdesk_tickets.json")
        if os.path.exists(file_path):
            with open(file_path, "r", encoding="utf-8") as f:
                tickets = json.load(f)
            for t in tickets:
                content = f"Customer Question:\n{t['question']}\n\nOfficial Answer / Resolution:\n{t['answer']}"
                self.domain_docs["freshdesk_qa"].append(Document(page_content=content, metadata={"source": "freshdesk", "ticket_id": t.get("ticket_id")}))

    # ==========================================
    # 🧠 2번 뇌: web_data (정책 & 정보 전용)
    # ==========================================
    def process_word_policies(self):
        print("🔍 [4/7] 정책 문서(.docx) 파싱 중... -> [web_data] 할당")
        docx_files = glob.glob(os.path.join(RAW_DIR, "*.docx")) # 💡 Warranty.docx, Sales Policy.docx 2개 동시 처리
        for file_path in docx_files:
            doc = docx.Document(file_path)
            content = " ".join([para.text.strip() for para in doc.paragraphs if para.text.strip()])
            self.domain_docs["web_data"].append(Document(page_content=content, metadata={"source": os.path.basename(file_path), "type": "policy"}))

    def process_web_data(self):
        print("🔍 [5/7] 웹사이트 크롤링 데이터 파싱 중... -> [web_data] 할당")
        file_path = os.path.join(DATA_DIR, "web_crawled_data.json")
        if os.path.exists(file_path):
            with open(file_path, "r", encoding="utf-8") as f:
                web_data = json.load(f)
            for data in web_data:
                self.domain_docs["web_data"].append(Document(page_content=data['content'], metadata={"source": data.get("source", "web")}))

    def process_curated_knowledge(self):
        print("🔍 [5.5/7] 큐레이션 지식 데이터 파싱 중... -> [web_data] 할당")
        file_path = os.path.join(DATA_DIR, "curated_knowledge.json")
        if os.path.exists(file_path):
            with open(file_path, "r", encoding="utf-8") as f:
                entries = json.load(f)
            for entry in entries:
                self.domain_docs["web_data"].append(Document(
                    page_content=entry['content'],
                    metadata={
                        "source": entry.get("source", "curated"),
                        "category": entry.get("category", "general"),
                        "type": "curated_knowledge"
                    }
                ))
            print(f"   ✅ 큐레이션 지식 {len(entries)}건 로드 완료")

    # ==========================================
    # 🧠 3번 뇌: osaki_products (상품 스펙 전용)
    # ==========================================
    def process_shopify_data(self):
        print("🔍 [6/7] 쇼피파이 상품 정제 데이터 파싱 중... -> [osaki_products] 할당")
        file_path = os.path.join(DATA_DIR, "cleaned_osaki_products.csv")
        if os.path.exists(file_path):
            df = pd.read_csv(file_path).fillna("")
            for _, row in df.iterrows():
                self.domain_docs["osaki_products"].append(Document(page_content=str(row.get('content', '')), metadata={"source": row.get('source', 'shopify')}))

    # 🌟 [수정완료] 7번째 파이프라인: Excel 멀티 레벨 헤더 정확 파싱 + 2단계 스마트 매칭
    def process_specifications(self):
        print("🔍 [7/7] 원본 쇼피파이 & 스펙 시트 융합(Join) 중... -> [osaki_products] 할당")
        spec_xlsx = os.path.join(RAW_DIR, "Specification_Massage Chair.xlsx")
        spec_csv  = os.path.join(RAW_DIR, "Specification_Massage Chair - Massage Chair.csv")
        shopify_csv = os.path.join(RAW_DIR, "products_export.csv")

        SKIP_COLS = {'join_key', 'Body (HTML)', 'Handle', 'Image Src', 'Image Position',
                     'Variant Image', 'Variant Grams', 'Variant Fulfillment Service',
                     'full_name_spec', 'Brand', 'Name', 'Manufacturer', 'Manufacturer Code',
                     'Update (YYYYMM)', 'Design Rep.', 'Video Rep.', 'Coder Rep.'}

        spec_source_exists = os.path.exists(spec_xlsx) or os.path.exists(spec_csv)
        if not (spec_source_exists and os.path.exists(shopify_csv)):
            print("   ⚠️  Spec file or Shopify CSV not found. Skipping.")
            return

        # 1. Excel 우선 — 멀티 레벨 헤더 (rows 1, 2, 3) 합쳐서 의미 보존
        if os.path.exists(spec_xlsx):
            print(f"   📊 Excel 파일 읽기: {os.path.basename(spec_xlsx)} (멀티 헤더)")
            # rows 0,1,2 모두 읽고 row 3부터 데이터
            raw = pd.read_excel(spec_xlsx, sheet_name="Massage Chair",
                                header=None, engine="openpyxl").fillna("")
            # 헤더 3개 행 (Excel index 1,2,3 → row 0 is empty, row 1=top group,
            # row 2=mid group, row 3=column header)
            top_row = raw.iloc[1].ffill()    # forward-fill across columns: "Dimension"...
            mid_row = raw.iloc[2].ffill()    # "Standing", "Minimum\nDoorway"...
            col_row = raw.iloc[3]            # "Width (inches)", "Assembled (inches)"...

            def clean(s):
                return str(s).replace("\r", " ").replace("\n", " ").strip()

            # 의미 있는 결합 컬럼명 만들기 — top + mid + col, 중복 제거
            combined_columns = []
            for t, m, c in zip(top_row, mid_row, col_row):
                t, m, c = clean(t), clean(m), clean(c)
                parts = []
                for p in (t, m, c):
                    if p and p not in parts and p not in ("Model", ""):
                        parts.append(p)
                combined_columns.append(" - ".join(parts) if parts else "")

            # 같은 이름 충돌 방지 (서로 다른 그룹 아래 같은 leaf 이름이 있을 수 있음)
            seen: dict = {}
            unique_columns = []
            for col in combined_columns:
                if not col:
                    unique_columns.append(f"_unnamed_{len(unique_columns)}")
                    continue
                if col in seen:
                    seen[col] += 1
                    unique_columns.append(f"{col} ({seen[col]})")
                else:
                    seen[col] = 1
                    unique_columns.append(col)

            # 데이터는 row 4 이후
            df_spec = raw.iloc[4:].copy()
            df_spec.columns = unique_columns
            df_spec = df_spec.replace("", "N/A").fillna("N/A")
            # 이름 없는 컬럼 제거
            df_spec = df_spec.loc[:, [c for c in df_spec.columns if not c.startswith("_unnamed_")]]
        else:
            print(f"   📄 CSV 파일 읽기 (Excel 없음): {os.path.basename(spec_csv)}")
            df_spec = pd.read_csv(spec_csv, skiprows=3, low_memory=False).fillna("N/A")

        df_spec.columns = [str(c).strip() for c in df_spec.columns]

        # Brand / Name 컬럼 찾기 (combined name has prefixes from top-row)
        def find_col(candidates):
            for c in df_spec.columns:
                if c.endswith(" - Brand") or c == "Brand":
                    return c if "Brand" in candidates else None
            return None
        # robust resolver
        brand_col = next((c for c in df_spec.columns if c.split(" - ")[-1].strip() == "Brand"), None)
        name_col  = next((c for c in df_spec.columns if c.split(" - ")[-1].strip() == "Name"), None)
        if brand_col and brand_col != "Brand":
            df_spec.rename(columns={brand_col: "Brand"}, inplace=True)
        if name_col and name_col != "Name":
            df_spec.rename(columns={name_col: "Name"}, inplace=True)

        # 빈 행 제거 (Brand 또는 Name이 없는 행)
        df_spec = df_spec[df_spec['Brand'].astype(str).str.strip().replace('N/A', '') != '']
        print(f"   총 {len(df_spec)}개 스펙 모델 로드됨.")
        # 디버그: 도어 관련 컬럼이 살아있는지 확인
        door_cols = [c for c in df_spec.columns if 'door' in c.lower() or 'minimum' in c.lower()]
        if door_cols:
            print(f"   🚪 도어 관련 컬럼 보존됨: {door_cols}")

        # 2. Shopify data (deduplicate by Title to avoid repeated rows per variant)
        df_shopify = pd.read_csv(shopify_csv, low_memory=False).fillna("N/A")
        df_shopify = df_shopify.drop_duplicates(subset=['Title'])

        # 3. Normalise for join
        # - Strips "massage chair", "os-" model prefix, 3D/4D/5D variant suffixes,
        #   punctuation and whitespace so titles like "Osaki Nova II 3D+" match
        #   the spec entry "Osaki Nova II".
        import re as _re
        def normalize_text(text):
            if pd.isna(text) or str(text) == "N/A": return ""
            s = str(text).lower()
            s = s.replace("massage chair", "")
            s = _re.sub(r'\bos-', '', s)          # strip "OS-" model prefix
            s = _re.sub(r'[^a-z0-9]', '', s)      # keep alphanumeric only
            return s.strip()

        # 4. Build full model name from Brand + Name
        df_spec['full_name_spec'] = (
            df_spec['Brand'].fillna('').astype(str) + ' ' +
            df_spec['Name'].fillna('').astype(str)
        ).str.strip()
        df_spec['join_key'] = df_spec['full_name_spec'].apply(normalize_text)
        df_spec['name_key'] = df_spec['Name'].fillna('').astype(str).apply(normalize_text)
        df_shopify['join_key'] = df_shopify['Title'].apply(normalize_text)

        # ── Phase 1: exact join ──────────────────────────────────────────
        merged_exact = pd.merge(df_shopify, df_spec, on='join_key', how='inner')
        matched_shop_keys = set(merged_exact['join_key'])
        print(f"   Phase 1 (exact):    {len(merged_exact)} rows matched.")

        # ── Phase 2: multi-strategy fuzzy fallback ──────────────────────
        # Strategy A: Brand+Name join_key bidirectional substring
        #   "Osaki Nova II 3D+"  ↔  "Osaki Nova II"
        #   "Titan Forge"        ↔  "Titan Forge 3D"
        # Strategy B: spec model NAME alone matches inside shop title
        #   "Osaki Hypnos 4D Pro AI" ⊃ "Hypnos"  (Brand="Osaki Platinum")
        # Strategy C: dimension-stripped substring
        unmatched_shop = df_shopify[~df_shopify['join_key'].isin(matched_shop_keys)]
        spec_keys = df_spec['join_key'].tolist()
        spec_name_keys = df_spec['name_key'].tolist()

        # noise words to ignore inside shop titles when matching name_key
        NOISE_WORDS = {"plus", "pro", "ai", "smart", "duo", "flex", "le", "lt", "ii", "iii", "xl", "xt"}

        def strip_dims(key):
            return _re.sub(r'\d+d\+?', '', key)

        def core_tokens(text):
            tokens = _re.findall(r"[a-z0-9]+", str(text).lower())
            return [t for t in tokens if t not in NOISE_WORDS and len(t) >= 3]

        fuzzy_rows = []
        matched_spec_idx = set()  # spec rows we successfully matched to a shop product
        for _, shop_row in unmatched_shop.iterrows():
            shop_key = shop_row['join_key']
            shop_title_lower = str(shop_row.get('Title', '')).lower()
            shop_vendor_lower = str(shop_row.get('Vendor', '')).lower()
            # Combine title AND vendor tokens — many products have brand in Vendor only
            # (e.g. Title="4D Orion Duo Mech", Vendor="Osaki").
            shop_tokens = core_tokens(shop_title_lower) + core_tokens(shop_vendor_lower)

            best_idx, best_score = None, 0

            for i, spec_key in enumerate(spec_keys):
                if not spec_key or len(spec_key) < 4:
                    continue
                # Strategy A: bidirectional substring of full join_key
                if (spec_key in shop_key or shop_key in spec_key) and len(spec_key) > best_score:
                    best_idx, best_score = i, len(spec_key)
                    continue
                # Strategy C: dimension-stripped substring
                sk_s, sh_s = strip_dims(spec_key), strip_dims(shop_key)
                if sk_s and len(sk_s) >= 6 and (sk_s in sh_s or sh_s in sk_s) and len(sk_s) > best_score:
                    best_idx, best_score = i, len(sk_s)
                    continue
                # Strategy B: name-only token match
                name_key = spec_name_keys[i]
                if not name_key or len(name_key) < 4:
                    continue
                if name_key in shop_key:
                    name_tokens = core_tokens(df_spec.iloc[i]['Name'])
                    spec_brand_tokens = core_tokens(df_spec.iloc[i]['Brand'])
                    # Require model NAME tokens to overlap; brand check is now lenient
                    # because shop_tokens already includes Vendor.
                    name_match = name_tokens and all(t in shop_tokens for t in name_tokens)
                    brand_match = (
                        not spec_brand_tokens  # spec has no brand → don't require
                        or any(b in shop_tokens for b in spec_brand_tokens)
                    )
                    if name_match and brand_match and len(name_key) > best_score:
                        best_idx, best_score = i, len(name_key)

            if best_idx is not None:
                fuzzy_rows.append({**shop_row.to_dict(), **df_spec.iloc[best_idx].to_dict()})
                matched_spec_idx.add(best_idx)

        fuzzy_df = pd.DataFrame(fuzzy_rows) if fuzzy_rows else pd.DataFrame()
        print(f"   Phase 2 (fuzzy):    {len(fuzzy_df)} additional rows matched.")

        merged_df = pd.concat([merged_exact, fuzzy_df], ignore_index=True)
        # Track which spec rows got matched (by join_key) so we can emit the
        # remaining ones as standalone spec docs in Phase 3.
        matched_spec_keys = set(merged_df['join_key'].dropna().astype(str).tolist())
        print(f"   ✅ Total spec join: {len(merged_df)} rows matched.")

        for _, row in merged_df.iterrows():
            model_name = str(row.get('Title', 'Unknown'))
            specs_text = []
            for col_name, value in row.items():
                if col_name in SKIP_COLS: continue
                val_str = str(value).strip()
                if not val_str or val_str in ("N/A", "nan"): continue
                specs_text.append(f"- {col_name}: {val_str}")

            content = f"Specifications for Model [{model_name}]:\n" + "\n".join(specs_text)
            self.domain_docs["osaki_products"].append(Document(
                page_content=content,
                metadata={"source": "specification_join", "title": model_name, "type": "specification"}
            ))

        # ── Phase 3: spec-only docs for rows that never matched a Shopify product ──
        # Without this, models like "Osaki Pro Soho II" (in spec sheet but missing
        # from products_export.csv) become invisible to the bot. Embedding them
        # standalone makes them searchable by name even with no shop record.
        unmatched_specs = df_spec[~df_spec['join_key'].isin(matched_spec_keys)]
        print(f"   Phase 3 (spec-only): {len(unmatched_specs)} unmatched spec rows emitted as standalone docs.")
        for _, spec_row in unmatched_specs.iterrows():
            full_name = str(spec_row.get('full_name_spec', '')).strip()
            if not full_name:
                continue
            specs_text = []
            for col_name, value in spec_row.items():
                if col_name in SKIP_COLS:
                    continue
                val_str = str(value).strip()
                if not val_str or val_str in ("N/A", "nan"):
                    continue
                specs_text.append(f"- {col_name}: {val_str}")
            if not specs_text:
                continue
            content = f"Specifications for Model [{full_name}]:\n" + "\n".join(specs_text)
            self.domain_docs["osaki_products"].append(Document(
                page_content=content,
                metadata={
                    "source": "specification_only",
                    "title": full_name,
                    "type": "specification",
                },
            ))

    # ==========================================
    # 🚀 다중 벡터 DB 동시 빌드 (Multi-Index Generation)
    # ==========================================
    def build_vector_dbs(self):
        print("\n🚀 [임베딩 시작] 도메인별 3개의 Vector DB를 생성합니다...")
        
        for domain_name, docs in self.domain_docs.items():
            if not docs:
                print(f"⚠️ [{domain_name}] 수집된 데이터가 없어 건너뜁니다.")
                continue
                
            print(f"🧠 [{domain_name}] 총 {len(docs)}개의 문서 임베딩 중...")
            vs = LC_FAISS.from_documents(docs, self.embeddings)
            
            save_path = os.path.join(FAISS_DIR, domain_name)
            vs.save_local(save_path)
            print(f"💾 [{domain_name}] DB 구축 완료! ({save_path})")

if __name__ == "__main__":
    os.makedirs(FAISS_DIR, exist_ok=True)
    
    ingester = MasterIngester()
    
    # [CS/Error 뇌세포]
    ingester.process_error_manuals()    # Auto-Check, fault_judgment
    ingester.process_qa_reports()       # Warranty Daily Report - Q&A
    ingester.process_freshdesk_tickets()# freshdesk_tickets.json
    
    # [Policy/Web 뇌세포]
    ingester.process_word_policies()    # Warranty.docx, Sales Policy.docx (glob으로 한 번에 2개 섭취)
    ingester.process_web_data()         # web_crawled_data.json
    ingester.process_curated_knowledge()# curated_knowledge.json (추천/FAQ/기능교육)
    
    # [Products/Specs 뇌세포]
    ingester.process_shopify_data()     # cleaned_osaki_products.csv
    ingester.process_specifications()   # products_export.csv + Specification_Massage Chair
    
    ingester.build_vector_dbs()
    print("\n🎉 모든 파이프라인이 성공적으로 완료되었습니다! 7개의 원본 파일과 3개의 정제 파일이 모두 3개의 뇌로 흡수되었습니다.")